from __future__ import annotations

import csv
import json
import os
import re
import subprocess
import tempfile
from email import policy
from email.parser import BytesParser
from html.parser import HTMLParser
from pathlib import Path

from app.config import Settings

TEXT_EXTENSIONS = {
    ".txt",
    ".text",
    ".md",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".xml",
    ".html",
    ".htm",
    ".eml",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
OFFICE_EXTENSIONS = {".docx", ".xlsx"}


def extract_text_from_file(path: Path, settings: Settings) -> tuple[str, str, list[str]]:
    warnings: list[str] = []
    suffix = path.suffix.lower()

    if not path.exists():
        return "", "missing_file", [f"File does not exist: {path}"]

    if suffix == ".pdf":
        return _extract_pdf(path, settings, warnings)

    if suffix in IMAGE_EXTENSIONS:
        return _extract_image(path, settings, warnings)

    if suffix in TEXT_EXTENSIONS:
        return _extract_text_like(path, suffix, warnings), _method_for_text_suffix(suffix), warnings

    if suffix in OFFICE_EXTENSIONS:
        return _extract_office(path, suffix, warnings)

    decoded = _decode_best_effort(path.read_bytes())
    if _looks_like_text(decoded):
        warnings.append(f"Unknown extension {suffix or '(none)'}; decoded as plain text.")
        return decoded, "plain_text_fallback", warnings

    warnings.append(f"Unsupported binary document extension: {suffix or '(none)'}")
    return "", "unsupported_binary", warnings


def _extract_pdf(path: Path, settings: Settings, warnings: list[str]) -> tuple[str, str, list[str]]:
    text = _extract_pdf_text(path, warnings)
    if text.strip():
        return text, "pdf_text", warnings

    if not settings.enable_ocr:
        warnings.append("PDF had no extractable text and OCR is disabled.")
        return "", "pdf_no_text", warnings

    ocr_text = _ocr_pdf_pages(path, settings, warnings)
    return ocr_text, "pdf_ocr", warnings


def _extract_image(path: Path, settings: Settings, warnings: list[str]) -> tuple[str, str, list[str]]:
    if not settings.enable_ocr:
        warnings.append("Image OCR is disabled.")
        return "", "image_ocr_disabled", warnings

    return _ocr_file(path, settings, warnings), "image_ocr", warnings


def _extract_text_like(path: Path, suffix: str, warnings: list[str]) -> str:
    if suffix == ".eml":
        return _extract_eml_text(path, warnings)

    text = _decode_best_effort(path.read_bytes())
    if suffix in {".html", ".htm", ".xml"}:
        return _strip_markup(text)
    if suffix in {".csv", ".tsv"}:
        return _extract_delimited_text(text, delimiter="\t" if suffix == ".tsv" else ",")
    if suffix in {".json", ".jsonl"}:
        return _extract_json_text(text, warnings)
    return text


def _extract_office(path: Path, suffix: str, warnings: list[str]) -> tuple[str, str, list[str]]:
    if suffix == ".docx":
        return _extract_docx(path, warnings), "docx_text", warnings
    if suffix == ".xlsx":
        return _extract_xlsx(path, warnings), "xlsx_text", warnings
    warnings.append(f"Unsupported Office extension: {suffix}")
    return "", "unsupported_office", warnings


def _extract_pdf_text(path: Path, warnings: list[str]) -> str:
    try:
        import pdfplumber
    except ImportError:
        warnings.append("pdfplumber is not installed; PDF text extraction skipped.")
        return ""

    try:
        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(f"--- page {index} ---\n{page_text}")
        return "\n\n".join(pages)
    except Exception as exc:
        warnings.append(f"PDF text extraction failed: {exc}")
        return ""


def _ocr_pdf_pages(path: Path, settings: Settings, warnings: list[str]) -> str:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        warnings.append("pypdfium2 is not installed; scanned PDF OCR skipped.")
        return ""

    texts: list[str] = []
    settings.cache_dir.mkdir(parents=True, exist_ok=True)
    try:
        pdf = pdfium.PdfDocument(str(path))
    except Exception as exc:
        warnings.append(f"Could not render PDF for OCR: {exc}")
        return ""

    with tempfile.TemporaryDirectory(dir=settings.cache_dir) as tmpdir:
        for index in range(len(pdf)):
            image_path = Path(tmpdir) / f"page-{index + 1}.png"
            try:
                bitmap = pdf[index].render(scale=2).to_pil()
                bitmap.save(image_path)
            except Exception as exc:
                warnings.append(f"Could not render PDF page {index + 1} for OCR: {exc}")
                continue
            page_text = _ocr_file(image_path, settings, warnings)
            if page_text.strip():
                texts.append(f"--- page {index + 1} OCR ---\n{page_text}")
    return "\n\n".join(texts)


def _ocr_file(path: Path, settings: Settings, warnings: list[str]) -> str:
    settings.cache_dir.mkdir(parents=True, exist_ok=True)
    engine = settings.ocr_engine.lower()

    if engine == "paddleocr":
        text = _paddle_ocr_file(path, settings, warnings)
        if text.strip():
            return text
        warnings.append("PaddleOCR unavailable or returned no text; falling back to Tesseract.")

    return _tesseract_ocr_file(path, settings, warnings)


def _paddle_ocr_file(path: Path, settings: Settings, warnings: list[str]) -> str:
    env = os.environ.copy()
    env.setdefault("PADDLE_PDX_CACHE_HOME", str(settings.cache_dir / "paddlex"))
    env.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
    script = """
import json
import sys
from paddleocr import PaddleOCR

def flatten(result):
    lines = []
    def walk(value):
        if isinstance(value, dict):
            for key in ("rec_texts", "texts"):
                texts = value.get(key)
                if isinstance(texts, list):
                    lines.extend(str(item) for item in texts if item)
            for item in value.values():
                walk(item)
            return
        if isinstance(value, (list, tuple)):
            if len(value) >= 2 and isinstance(value[1], (list, tuple)) and value[1]:
                candidate = value[1][0]
                if isinstance(candidate, str):
                    lines.append(candidate)
                    return
            for item in value:
                walk(item)
    walk(result)
    return "\\n".join(dict.fromkeys(line.strip() for line in lines if line.strip()))

ocr = PaddleOCR(lang="en")
print(flatten(ocr.ocr(sys.argv[1])))
"""
    try:
        result = subprocess.run(
            [settings_python_executable(), "-c", script, str(path)],
            check=False,
            capture_output=True,
            env=env,
            text=True,
            timeout=settings.paddle_ocr_timeout_seconds,
        )
    except subprocess.TimeoutExpired:
        warnings.append(f"PaddleOCR timed out after {settings.paddle_ocr_timeout_seconds}s.")
        return ""
    except Exception as exc:
        warnings.append(f"PaddleOCR could not start: {exc}")
        return ""

    if result.returncode != 0:
        warnings.append(_compact_warning("PaddleOCR failed", result.stderr))
        return ""
    return result.stdout.strip()


def _tesseract_ocr_file(path: Path, settings: Settings, warnings: list[str]) -> str:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        warnings.append("pytesseract/Pillow not installed; OCR skipped.")
        return ""

    if settings.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    try:
        return pytesseract.image_to_string(Image.open(path))
    except Exception as exc:
        warnings.append(f"Tesseract OCR failed: {exc}")
        return ""


def _extract_eml_text(path: Path, warnings: list[str]) -> str:
    try:
        message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    except Exception as exc:
        warnings.append(f"Email parse failed: {exc}")
        return ""

    chunks: list[str] = []
    for header in ("From", "To", "Subject", "Date"):
        value = message.get(header)
        if value:
            chunks.append(f"{header}: {value}")

    for part in message.walk():
        if part.get_content_disposition() == "attachment":
            filename = part.get_filename() or "attachment"
            chunks.append(f"Attachment: {filename} ({part.get_content_type()})")
            continue
        content_type = part.get_content_type()
        if content_type in {"text/plain", "text/html"}:
            try:
                payload = part.get_content()
            except Exception:
                continue
            chunks.append(_strip_markup(payload) if content_type == "text/html" else payload)

    return "\n\n".join(chunk.strip() for chunk in chunks if chunk and chunk.strip())


def _extract_docx(path: Path, warnings: list[str]) -> str:
    try:
        from docx import Document
    except ImportError:
        warnings.append("python-docx is not installed; DOCX extraction skipped.")
        return ""

    try:
        document = Document(path)
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    except Exception as exc:
        warnings.append(f"DOCX extraction failed: {exc}")
        return ""


def _extract_xlsx(path: Path, warnings: list[str]) -> str:
    try:
        import openpyxl
    except ImportError:
        warnings.append("openpyxl is not installed; XLSX extraction skipped.")
        return ""

    try:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        rows: list[str] = []
        for sheet in workbook.worksheets:
            rows.append(f"--- sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    rows.append("\t".join(values))
        return "\n".join(rows)
    except Exception as exc:
        warnings.append(f"XLSX extraction failed: {exc}")
        return ""


def _extract_json_text(text: str, warnings: list[str]) -> str:
    try:
        parsed = json.loads(text)
        return json.dumps(parsed, indent=2, sort_keys=True)
    except json.JSONDecodeError:
        warnings.append("JSON parse failed; using raw decoded text.")
        return text


def _extract_delimited_text(text: str, delimiter: str) -> str:
    rows = csv.reader(text.splitlines(), delimiter=delimiter)
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)


class _MarkupStripper(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.parts.append(data.strip())


def _strip_markup(text: str) -> str:
    parser = _MarkupStripper()
    parser.feed(text)
    stripped = "\n".join(parser.parts)
    return stripped or re.sub(r"<[^>]+>", " ", text)


def _decode_best_effort(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _looks_like_text(text: str) -> bool:
    if not text:
        return False
    printable = sum(character.isprintable() or character.isspace() for character in text)
    return printable / max(len(text), 1) > 0.85


def _method_for_text_suffix(suffix: str) -> str:
    return {
        ".csv": "csv_text",
        ".tsv": "tsv_text",
        ".json": "json_text",
        ".jsonl": "json_text",
        ".xml": "markup_text",
        ".html": "markup_text",
        ".htm": "markup_text",
        ".eml": "email_text",
    }.get(suffix, "plain_text")


def _compact_warning(prefix: str, value: str) -> str:
    one_line = " ".join(value.split())
    return f"{prefix}: {one_line[:400]}"


def settings_python_executable() -> str:
    import sys

    return sys.executable
